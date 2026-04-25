"""Gera a proposta em formato .docx a partir do conte�do estruturado.

Uso:
    venv\\Scripts\\python.exe docs\\proposta\\generate_docx.py
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path

ACCENT = RGBColor(0x0A, 0x7D, 0x4E)
MUTED = RGBColor(0x6B, 0x6B, 0x6B)
DARK = RGBColor(0x1A, 0x1A, 0x1A)


def set_cell_bg(cell, color_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(14)
        run.font.color.rgb = ACCENT
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(6)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "E2E2E2")
        pBdr.append(bottom)
        pPr.append(pBdr)
    elif level == 2:
        run.font.size = Pt(11)
        run.font.color.rgb = DARK
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)


def add_para(doc: Document, text: str, *, bold: bool = False, color=None, size: int = 10):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    if bold:
        run.bold = True
    if color:
        run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(6)
    return p


def add_bullets(doc: Document, items: list[str]) -> None:
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.runs[0] if p.runs else p.add_run("")
        run.text = it
        run.font.size = Pt(10)


def add_two_column_meta(doc: Document, pairs: list[tuple[str, str]]) -> None:
    rows = (len(pairs) + 1) // 2
    table = doc.add_table(rows=rows, cols=4)
    table.autofit = True
    for i, (label, value) in enumerate(pairs):
        row = i // 2
        col = (i % 2) * 2
        lcell = table.cell(row, col)
        vcell = table.cell(row, col + 1)
        lcell.text = ""
        vcell.text = ""
        lp = lcell.paragraphs[0]
        lr = lp.add_run(label)
        lr.bold = True
        lr.font.size = Pt(9)
        lr.font.color.rgb = MUTED
        vp = vcell.paragraphs[0]
        vr = vp.add_run(value)
        vr.font.size = Pt(9)
        set_cell_bg(lcell, "FAFAFA")
        set_cell_bg(vcell, "FAFAFA")


def add_simple_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = ACCENT
        set_cell_bg(hdr[i], "E8F5EE")
    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            cell = table.cell(ri, ci)
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(val)
            run.font.size = Pt(9)


def add_price_box(doc: Document) -> None:
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    set_cell_bg(cell, "E8F5EE")
    cell.text = ""

    p1 = cell.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p1.add_run("VALOR TOTAL DO PROJETO")
    r1.font.size = Pt(9)
    r1.font.color.rgb = MUTED
    r1.bold = True

    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("R$ 6.500,00")
    r2.font.size = Pt(28)
    r2.bold = True
    r2.font.color.rgb = ACCENT

    p3 = cell.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run("Licen�a perp�tua + implanta��o + treinamento + 30 dias de garantia")
    r3.font.size = Pt(9)
    r3.font.color.rgb = MUTED


def build() -> Document:
    doc = Document()

    # Margens
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    # Estilo padr�o
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    # Header / capa
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    ribbon = p.add_run("PROPOSTA COMERCIAL")
    ribbon.bold = True
    ribbon.font.size = Pt(9)
    ribbon.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    # ribbon shading via paragraph background n�o � trivial; usamos cor de fonte sobre destaque visual
    p2 = doc.add_paragraph()
    title = p2.add_run("Photo Manager")
    title.bold = True
    title.font.size = Pt(24)
    p3 = doc.add_paragraph()
    sub = p3.add_run("Sistema de gerenciamento e reconhecimento facial para fotografia de eventos")
    sub.font.size = Pt(11)
    sub.font.color.rgb = MUTED

    add_two_column_meta(doc, [
        ("Cliente:", "Eduardo Santana Produ��es"),
        ("Fornecedor:", "Bruno Dias"),
        ("Refer�ncia:", "PROP-2026-04-001"),
        ("Data:", "25 de abril de 2026"),
        ("Validade:", "30 dias"),
        ("CPF:", "414.078.848-83"),
    ])

    # 1. Apresenta��o
    add_heading(doc, "1. Apresenta��o", 1)
    add_para(doc, "Esta proposta tem como objetivo formalizar o fornecimento e a implanta��o do Photo Manager, um software desktop de gerenciamento e processamento automatizado de fotos para eventos, com tecnologia de reconhecimento facial.")
    add_para(doc, "A solu��o foi desenvolvida sob medida para fluxos de trabalho de fotografia em eventos, permitindo que o operador capture rostos de refer�ncia via webcam, escaneie sess�es de fotos e identifique automaticamente todas as imagens em que cada pessoa aparece � eliminando a necessidade de revis�o manual foto a foto.")

    # 2. Descri��o
    add_heading(doc, "2. Descri��o da Solu��o", 1)
    add_para(doc, "O Photo Manager � uma aplica��o local executada no pr�prio computador do cliente (sem necessidade de internet ou servidor externo) com interface moderna acess�vel pelo navegador. O fluxo � dividido em quatro etapas:")
    for i, (titulo, desc) in enumerate([
        ("Sess�o", "sele��o da pasta com as fotos do evento"),
        ("Captura", "registro dos rostos de refer�ncia via webcam"),
        ("Fotos", "escaneamento automatizado e exibi��o das imagens correspondentes"),
        ("Imprimir", "composi��o final com template personalizado e envio para impressora"),
    ], start=1):
        p = doc.add_paragraph(style="List Number")
        r1 = p.add_run(f"{titulo}")
        r1.bold = True
        r1.font.size = Pt(10)
        r2 = p.add_run(f" � {desc}")
        r2.font.size = Pt(10)

    # 3. Funcionalidades
    add_heading(doc, "3. Funcionalidades Entregues", 1)

    add_heading(doc, "Captura e processamento de rostos", 2)
    add_bullets(doc, [
        "Captura de rostos via webcam com detec��o autom�tica",
        "Suporte a m�ltiplos rostos de refer�ncia simult�neos",
        "Algoritmo de reconhecimento facial baseado em dlib e face_recognition (estado da arte)",
        "Cache de detec��es por imagem (rescans subsequentes em segundos)",
        "Modos de scan ajust�veis: r�pido (padr�o) e preciso (rostos pequenos/escuros)",
    ])

    add_heading(doc, "Gerenciamento de sess�es", 2)
    add_bullets(doc, [
        "Sele��o de pasta de fotos via Explorador de Arquivos nativo do Windows",
        "Hist�rico das 10 sess�es mais recentes com indicadores de quantidade e tempo de processamento",
        "Suporte a formatos JPG, JPEG e PNG",
        "Estat�sticas de progresso em tempo real durante o scan",
    ])

    add_heading(doc, "Composi��o e impress�o", 2)
    add_bullets(doc, [
        "Templates de composi��o personalizados (footer com logo/marca do evento)",
        "Modos de enquadramento: cobrir (cover) ou conter (contain)",
        "Alinhamento vertical configur�vel (topo, centro, rodap�)",
        "Suporte a orienta��o retrato e paisagem por foto",
        "Envio direto para impressora padr�o do sistema",
    ])

    add_heading(doc, "Interface", 2)
    add_bullets(doc, [
        "Tema dark e light com toggle vis�vel",
        "Carrossel de visualiza��o ampliada",
        "Barra de status com contagem de fotos e matches em tempo real",
        "Snackbar de notifica��es",
        "Layout responsivo otimizado para uso em desktop",
    ])

    add_heading(doc, "Performance e confiabilidade", 2)
    add_bullets(doc, [
        "Cache persistente de encodings em disco (rescans n�o recomputam)",
        "Progress incremental (crash no meio do scan n�o perde o trabalho j� feito)",
        "Logs detalhados em pasta dedicada para diagn�stico",
        "Thumbnails pr�-gerados por sess�o para navega��o fluida",
        "Suporte opcional a processamento multi-thread (configur�vel)",
    ])

    # 4. Stack
    add_heading(doc, "4. Stack Tecnol�gica", 1)
    add_simple_table(doc,
        ["Componente", "Tecnologia"],
        [
            ["Backend", "Python 3.10+ com FastAPI e Uvicorn"],
            ["Reconhecimento facial", "dlib + face_recognition"],
            ["Processamento de imagem", "OpenCV, Pillow, NumPy"],
            ["Frontend", "HTML5, CSS3, JavaScript ES2020 (sem depend�ncias externas)"],
            ["C�mera no navegador", "MediaPipe Vision"],
            ["Impress�o", "pywin32 (integra��o nativa Windows)"],
            ["Servidor", "Local (127.0.0.1), porta 8000"],
        ]
    )

    # 5. Requisitos
    add_heading(doc, "5. Requisitos de Hardware e Sistema", 1)
    add_heading(doc, "M�nimos", 2)
    add_bullets(doc, [
        "Sistema operacional: Windows 10 (64-bit) ou superior",
        "Processador: Intel Core i5 (8� gera��o) ou equivalente AMD Ryzen 5",
        "Mem�ria RAM: 8 GB",
        "Armazenamento: 2 GB livres + espa�o adicional para fotos",
        "Webcam: integrada ou USB (resolu��o m�nima 720p)",
        "Impressora: instalada e configurada como padr�o no Windows",
        "Navegador: Google Chrome, Microsoft Edge ou Firefox (vers�o atualizada)",
    ])
    add_heading(doc, "Recomendados (acima de 500 fotos por evento)", 2)
    add_bullets(doc, [
        "Processador: Intel Core i7 (10� gera��o ou superior) / AMD Ryzen 7",
        "Mem�ria RAM: 16 GB",
        "Armazenamento: SSD com 10 GB livres",
        "Webcam: Full HD (1080p) com boa ilumina��o no ambiente",
    ])

    # 6. Prazo
    add_heading(doc, "6. Prazo de Entrega e Instala��o", 1)
    add_simple_table(doc,
        ["Etapa", "Prazo"],
        [
            ["Assinatura da proposta e pagamento da entrada", "Dia 0"],
            ["Agendamento da instala��o", "At� 3 dias �teis ap�s entrada"],
            ["Instala��o presencial ou remota e configura��o", "1 dia �til"],
            ["Treinamento operacional (at� 2 horas)", "Mesmo dia da instala��o"],
            ["In�cio da garantia de 30 dias", "Dia da entrega"],
        ]
    )
    add_para(doc, "Total: at� 5 dias �teis entre assinatura e opera��o.", bold=True)

    # 7. Treinamento
    add_heading(doc, "7. Treinamento Incluso", 1)
    add_para(doc, "Est� incluso no valor da implanta��o uma sess�o de treinamento de at� 2 horas, presencial ou via videoconfer�ncia, cobrindo:")
    add_bullets(doc, [
        "Opera��o do fluxo completo (sess�o ? captura ? fotos ? impress�o)",
        "Boas pr�ticas para qualidade de reconhecimento (ilumina��o, dist�ncia, �ngulo)",
        "Gerenciamento de pastas de eventos e organiza��o de arquivos",
        "Uso dos templates de composi��o",
        "Configura��o da impressora e troubleshooting b�sico",
        "Diagn�stico de problemas via logs",
    ])
    add_para(doc, "Material de apoio em PDF � entregue ao final da sess�o.")

    # 8. Atualiza��es
    add_heading(doc, "8. Pol�tica de Atualiza��es", 1)
    add_bullets(doc, [
        "Corre��es de bugs durante a garantia (30 dias): sem custo adicional",
        "Corre��es de bugs ap�s a garantia: sob demanda, cobradas � parte por hora trabalhada",
        "Novas funcionalidades: or�amento separado, baseado em escopo e estimativa de horas",
        "Atualiza��es de bibliotecas/seguran�a: podem ser solicitadas a qualquer momento, cobradas como manuten��o sob demanda",
    ])

    # 9. Termos legais
    add_heading(doc, "9. Termos Legais", 1)

    add_heading(doc, "Propriedade intelectual", 2)
    add_para(doc, "O c�digo-fonte da aplica��o permanece de propriedade do fornecedor. O cliente recebe licen�a de uso perp�tua, n�o exclusiva e intransfer�vel para uso comercial nos seus pr�prios eventos.")

    add_heading(doc, "Confidencialidade (NDA)", 2)
    add_para(doc, "O fornecedor compromete-se a tratar como confidencial qualquer informa��o, lista de clientes, fotos ou material de eventos a que tenha acesso durante a instala��o, treinamento ou suporte. O cliente compromete-se a n�o redistribuir, sublicenciar ou compartilhar a aplica��o com terceiros.")

    add_heading(doc, "LGPD (Lei Geral de Prote��o de Dados)", 2)
    add_para(doc, "A aplica��o processa imagens faciais, classificadas como dados pessoais sens�veis pela LGPD (Art. 5�, II). � responsabilidade exclusiva do cliente (controlador dos dados):")
    add_bullets(doc, [
        "Obter consentimento expresso dos titulares (pessoas fotografadas) para captura, processamento e armazenamento das imagens",
        "Manter registros de consentimento conforme exigido pela ANPD",
        "Definir e cumprir prazo de reten��o das fotos",
        "Atender solicita��es de exclus�o por parte dos titulares",
        "Implementar medidas de seguran�a no ambiente onde a aplica��o opera",
    ])
    add_para(doc, "A aplica��o roda inteiramente local (sem envio de dados para nuvem ou terceiros), o que facilita o cumprimento da LGPD, mas a responsabilidade legal permanece do cliente.")

    add_heading(doc, "Garantia", 2)
    add_para(doc, "A garantia de 30 dias cobre exclusivamente corre��es de bugs e defeitos de funcionamento. N�o cobre:")
    add_bullets(doc, [
        "Falhas decorrentes de altera��o n�o autorizada do c�digo ou da configura��o",
        "Problemas causados por hardware defeituoso, drivers desatualizados ou ambiente de rede",
        "Mudan�as no sistema operacional ou em depend�ncias externas (impressoras, webcams) ap�s a entrega",
    ])

    add_heading(doc, "Limita��o de responsabilidade", 2)
    add_para(doc, "A responsabilidade total do fornecedor est� limitada ao valor pago pelo cliente nesta proposta. O fornecedor n�o responde por lucros cessantes, perdas indiretas ou danos a terceiros decorrentes do uso da aplica��o.")

    add_heading(doc, "Foro", 2)
    add_para(doc, "Fica eleito o foro da Comarca de S�o Paulo/SP para dirimir quaisquer quest�es oriundas deste contrato, com ren�ncia expressa a qualquer outro, por mais privilegiado que seja.")

    # 10. Investimento
    add_heading(doc, "10. Investimento", 1)
    add_price_box(doc)
    doc.add_paragraph()
    add_heading(doc, "Forma de pagamento", 2)
    add_simple_table(doc,
        ["Op��o", "Valor", "Detalhes"],
        [
            ["� vista", "R$ 6.175,00", "Com 5% de desconto, no ato da assinatura"],
            ["Parcelado em 2x", "2 � R$ 3.250,00", "50% na assinatura + 50% na entrega"],
        ]
    )
    add_para(doc, "Meios aceitos: PIX, transfer�ncia banc�ria ou boleto.", bold=True)

    add_heading(doc, "Ap�s o per�odo de garantia (a partir do 31� dia)", 2)
    add_simple_table(doc,
        ["Servi�o", "Valor"],
        [
            ["Corre��es de bugs sob demanda", "R$ 150,00 / hora"],
            ["Pequenos ajustes e configura��es", "R$ 150,00 / hora"],
            ["Novas funcionalidades", "Or�amento � parte mediante escopo"],
        ]
    )
    add_para(doc, "M�nimo de cobran�a: 1 hora por chamado.", color=MUTED, size=9)

    # 11. Aceite
    add_heading(doc, "11. Aceite", 1)
    add_para(doc, "Para aceitar esta proposta, basta responder a este documento confirmando os termos e o modelo de pagamento escolhido.")
    add_para(doc, "A partir da confirma��o, o fornecedor enviar� os dados banc�rios e iniciar� o agendamento da instala��o.")

    # Assinatura
    doc.add_paragraph()
    sig_p = doc.add_paragraph()
    sig_p.paragraph_format.space_before = Pt(24)
    pPr = sig_p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), "8")
    top.set(qn("w:space"), "1")
    top.set(qn("w:color"), "E2E2E2")
    pBdr.append(top)
    pPr.append(pBdr)
    name = sig_p.add_run("Bruno Dias")
    name.bold = True
    name.font.size = Pt(13)
    for line in [
        "CPF: 414.078.848-83 � S�o Paulo / SP",
        "Telefone / WhatsApp: +55 (11) 95448-5244",
        "E-mail: dev.brunodias@gmail.com",
    ]:
        p = doc.add_paragraph()
        r = p.add_run(line)
        r.font.size = Pt(9)
        r.font.color.rgb = MUTED
        p.paragraph_format.space_after = Pt(2)

    # Footer text
    doc.add_paragraph()
    foot = doc.add_paragraph()
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = foot.add_run("Proposta PROP-2026-04-001 emitida em 25/04/2026 � Validade: 30 dias.\nAp�s esse prazo, valores e condi��es est�o sujeitos a revis�o.")
    fr.font.size = Pt(8)
    fr.font.color.rgb = MUTED
    fr.italic = True

    return doc


if __name__ == "__main__":
    out = Path(__file__).parent / "proposta-eduardo-santana-producoes.docx"
    doc = build()
    doc.save(out)
    print(f"OK: {out}")
